from docx import Document
from docx.oxml.ns import qn
import pandas as pd
from pathlib import Path


def table_to_markdown(table):
    """Convert a python-docx Table to markdown."""
    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]

    if not rows:
        return ''

    # De-duplicate merged cells (docx repeats merged cell text)
    cleaned = []
    for row in rows:
        deduped = [row[0]]
        for i in range(1, len(row)):
            deduped.append('' if row[i] == row[i-1] else row[i])
        cleaned.append(deduped)

    header = cleaned[0]
    data   = cleaned[1:]
    n      = len(header)
    pad    = lambda r: r + [''] * (n - len(r))

    lines = ['| ' + ' | '.join(header) + ' |',
             '| ' + ' | '.join(['---'] * n) + ' |']
    for row in data:
        lines.append('| ' + ' | '.join(pad(row)) + ' |')
    return '\n'.join(lines)


def heading_level_to_md(style_name):
    """Map Word heading style to markdown # level.
    Handles both 'Heading 1' (para.style.name) and 'Heading1' (XML w:val)."""
    import re
    name = style_name.lower().replace(' ', '')
    if 'title' in name:
        return 1
    m = re.search(r'heading(\d)', name)
    if m:
        return int(m.group(1)) + 1   # Heading1=##, Heading2=###
    return None


def extract_docx_to_markdown(file_path: str) -> dict:
    path = Path(file_path)
    warnings = []
    parts = []

    try:
        doc = Document(file_path)

        # Track which paragraphs are inside tables (to avoid double-processing)
        table_paragraphs = set()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        table_paragraphs.add(id(para))

        # Process document body in order (paragraphs and tables)
        # doc.element.body contains both <w:p> and <w:tbl> elements
        table_index = 0

        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                # Find the matching paragraph object
                para_text = ''.join(
                    node.text for node in child.iter()
                    if node.tag.endswith('}t') and node.text
                ).strip()

                if not para_text:
                    continue

                # Get style name from XML
                style_node = child.find('.//' + qn('w:pStyle'))
                style_name = style_node.get(qn('w:val'), '') if style_node is not None else 'Normal'

                # Check if it's a list bullet
                num_node = child.find('.//' + qn('w:numPr'))
                is_list = num_node is not None or 'list' in style_name.lower() or 'bullet' in style_name.lower()

                # Map heading level
                level = heading_level_to_md(style_name)

                if level:
                    parts.append(f"\n{'#' * level} {para_text}\n")
                elif is_list:
                    parts.append(f"- {para_text}")
                else:
                    parts.append(para_text)

            elif tag == 'tbl':
                if table_index < len(doc.tables):
                    md = table_to_markdown(doc.tables[table_index])
                    if md:
                        parts.append(md)
                    table_index += 1

        full_md = '\n\n'.join(p for p in parts if p and p.strip())
        quality = ('failed'  if not full_md.strip() else
                   'partial' if warnings else
                   'full')

        return {
            'file'    : path.name,
            'markdown': full_md,
            'quality' : quality,
            'warnings': warnings
        }

    except Exception as e:
        return {
            'file'    : path.name,
            'markdown': '',
            'quality' : 'failed',
            'warnings': [f'Cannot open file: {e}']
        }


if __name__ == '__main__':
    result = extract_docx_to_markdown('/home/claude/test_completion_report.docx')
    print(f"File    : {result['file']}")
    print(f"Quality : {result['quality']}")
    if result['warnings']:
        print(f"Warnings: {result['warnings']}")
    print()
    print('=' * 60)
    print(result['markdown'])
